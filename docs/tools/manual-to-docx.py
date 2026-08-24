#!/usr/bin/env python3
"""
แปลงคู่มือผู้ทดสอบ (Markdown) เป็น .docx ที่จัดหน้าเรียบร้อยด้วยฟอนต์ Sarabun

    python3 docs/tools/manual-to-docx.py

อ่านจาก docs/1[2-5]-tester-manual-*.md กับ docs/18-document-template-variables.md
แล้วเขียนออกเป็น docs/manuals-docx/*.docx

ทำไมต้องเขียนเอง: เครื่องนี้ไม่มี pandoc และต่อให้มี ก็ยังต้องตั้งค่าฟอนต์ฝั่ง
complex script (w:cs / w:szCs) เองอยู่ดี — ถ้าไม่ตั้ง Word จะเรนเดอร์ภาษาไทย
ด้วยฟอนต์และขนาดของ default ไม่ใช่ Sarabun ตามที่สั่ง

ฟอนต์ถูก **ฝังลงไฟล์** (ECMA-376 §15.2.13 obfuscated font) คนเปิดจึงเห็น Sarabun
แม้เครื่องปลายทางจะไม่ได้ติดตั้งไว้
"""
import os
import re
import subprocess
import sys
import uuid
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
DOCS = ROOT / "docs"
OUT = DOCS / "manuals-docx"
FONT_ZIP = ROOT.parent / "assets" / "theme_ci_design" / "Font" / "Sarabun.zip"
LOGO = ROOT / "backend" / "src" / "assets" / "brand" / "bdi-logo.png"

# สีจาก CI ของ BDI — ชุดเดียวกับที่หน้าเว็บและ PDF ใช้
NAVY = RGBColor(0x19, 0x27, 0x68)
CORAL = RGBColor(0xE5, 0x77, 0x5A)
INK = RGBColor(0x1A, 0x1E, 0x33)
MUTED = RGBColor(0x5C, 0x63, 0x7A)
RULE = "D8DCE6"
BAND = "F2F4F9"
CODEBG = "F5F6FA"

BODY = "Sarabun"
MONO = "DejaVu Sans Mono"
BODY_PT = 11

# ไล่ขนาดหัวข้อจากเนื้อความขึ้นไปเป็นลำดับ — `##` ใหญ่สุด แล้วลดหลั่นลงมา
HEADING_PT = {1: 18, 2: 14, 3: 12}
TABLE_PT = 10
CAPTION_PT = 9.5
CODE_PT = 8.5
CHROME_PT = 9  # หัวกระดาษ/ท้ายกระดาษ

CONTENT_CM = 21.0 - 2.2 - 2.0  # A4 กว้าง 21 ซม. ลบขอบซ้าย/ขวา

TESTER = "คู่มือสำหรับผู้ทดสอบระบบ · ไม่ต้องมีพื้นฐานทางเทคนิค"

# (ไฟล์ Markdown, ชื่อบนปก, คำโปรย, บรรทัดบอกว่าเขียนให้ใคร, วันที่ปรับปรุง)
# วันที่เป็นของเนื้อหา ไม่ใช่ของวันที่รันสคริปต์ — สร้างไฟล์ใหม่โดยไม่ได้แก้เนื้อหา
# ไม่ควรทำให้ปกบอกว่าคู่มือถูกปรับปรุง
BOOKS = [
    ("12-tester-manual-overview.md", "คู่มือผู้ทดสอบระบบ", "ภาพรวมและการเข้าใช้งาน", TESTER, "18 สิงหาคม 2569"),
    ("13-tester-manual-journey-a.md", "เล่ม A", "ผู้ดูแลระบบสร้างหน่วยงานและส่งคำเชิญ", TESTER, "18 สิงหาคม 2569"),
    ("14-tester-manual-journey-b.md", "เล่ม B", "หน่วยงานลงทะเบียนตัวเองจนได้รับอนุมัติ", TESTER, "18 สิงหาคม 2569"),
    ("15-tester-manual-journey-c.md", "เล่ม C", "ลงทะเบียนชุดข้อมูลจนได้รับอนุมัติ", TESTER, "18 สิงหาคม 2569"),
    # ไม่ใช่คู่มือผู้ทดสอบ แต่เป็นคู่มือของผู้เขียนเอกสารต้นแบบ — เขาทำงานใน Word อยู่แล้ว
    # จึงต้องมีฉบับ .docx ให้เปิดคู่กันไปกับ template ที่กำลังแก้ (การ์ด Enhance ข้อ 2)
    (
        "18-document-template-variables.md",
        "ตัวแปรในเอกสารต้นแบบ",
        "รายการ placeholder ที่ใช้ได้ทั้งหมด",
        "คู่มือสำหรับผู้เขียนเอกสารต้นแบบ · ไม่ต้องเขียนโค้ด",
        "24 สิงหาคม 2569",
    ),
]

# ---------------------------------------------------------------- ตัวช่วยระดับ XML


def _set(el, tag, **attrs):
    child = OxmlElement(tag)
    for k, v in attrs.items():
        child.set(qn(k), str(v))
    el.append(child)
    return child


def style_run(run, *, font=BODY, size=BODY_PT, bold=False, italic=False, color=None):
    """
    ตั้งฟอนต์ให้ครบทั้งสามชุด — ascii / hAnsi / cs

    ภาษาไทยนับเป็น complex script ถ้าตั้งแต่ `w:sz` โดยไม่ตั้ง `w:szCs` ด้วย
    Word จะใช้ขนาดตั้งต้นกับตัวอักษรไทยทั้งหมด ผลคือเอกสารมีสองขนาดปนกัน
    """
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), font)
    for tag, val in (("w:sz", size), ("w:szCs", size)):
        _set(rPr, tag, **{"w:val": int(val * 2)})
    if bold:
        _set(rPr, "w:b", **{"w:val": "1"})
        _set(rPr, "w:bCs", **{"w:val": "1"})
    if italic:
        _set(rPr, "w:i", **{"w:val": "1"})
        _set(rPr, "w:iCs", **{"w:val": "1"})
    if color is not None:
        run.font.color.rgb = color
    return run


def shade(paragraph, fill):
    _set(paragraph._p.get_or_add_pPr(), "w:shd", **{"w:val": "clear", "w:fill": fill})


def borders(paragraph, *, left=None, box=None, bottom=None):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    if box:
        for side in ("top", "left", "bottom", "right"):
            _set(pBdr, f"w:{side}", **{"w:val": "single", "w:sz": "6", "w:space": "6", "w:color": box})
    if left:
        _set(pBdr, "w:left", **{"w:val": "single", "w:sz": "24", "w:space": "10", "w:color": left})
    if bottom:
        _set(pBdr, "w:bottom", **{"w:val": "single", "w:sz": "6", "w:space": "4", "w:color": bottom})
    pPr.append(pBdr)


def keep_with_next(paragraph):
    _set(paragraph._p.get_or_add_pPr(), "w:keepNext", **{"w:val": "1"})


def spacing(paragraph, *, before=0, after=6, line=1.45):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line


# ---------------------------------------------------------------- inline markdown

INLINE = re.compile(
    r"(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\)|<https?://[^>\s]+>|\*[^*\n]+\*)",
    re.S,
)


# อีโมจิสีที่หน้าเว็บใช้ได้ ไม่มีในฟอนต์เอกสาร — เปลี่ยนเป็นสัญลักษณ์ขาวดำที่มีจริง
EMOJI = {"✅": "✔", "❌": "✘", "⚠️": "⚠", "🎉": "", "⚠": "⚠"}
SYMBOL_FONT = "DejaVu Sans"
SYMBOLS = set("✔✘⚠☐▸←→↓")
SYMBOL_SPLIT = re.compile("([" + "".join(SYMBOLS) + "])")


def unescape(text):
    text = text.replace("&lt;", "<").replace("&gt;", ">").replace("&amp;", "&")
    for src, dst in EMOJI.items():
        text = text.replace(src, dst)
    return text.replace("  ", " ").strip() if text.strip() == "" else text


def add_text(paragraph, text, **kw):
    """
    เขียนข้อความลงย่อหน้า โดยแยกสัญลักษณ์ออกไปใช้ฟอนต์ที่มีตัวนั้นจริง

    Sarabun ไม่มี ✔ ⚠ ☐ ▸ ถ้าปล่อยไว้ในรันเดียวกัน ตัวที่หายไปจะถูกตัดทิ้งเงียบ ๆ
    (LibreOffice ทำแบบนั้นกับ ⚠️ มาแล้ว) ผู้อ่านจึงไม่เห็นเครื่องหมายเตือนเลย
    """
    for chunk in SYMBOL_SPLIT.split(text):
        if not chunk:
            continue
        if chunk in SYMBOLS:
            style_run(paragraph.add_run(chunk), **{**kw, "font": SYMBOL_FONT})
        else:
            style_run(paragraph.add_run(chunk), **kw)


def add_inline(paragraph, text, *, size=BODY_PT, color=INK, bold=False, italic=False):
    """
    เขียนข้อความหนึ่งย่อหน้า โดยแปลง **ตัวหนา** · `โค้ด` · *เอียง* · [ลิงก์](ปลายทาง)

    ตัวหนากับตัวเอียง **เรียกตัวเองซ้ำ** กับเนื้อข้างใน เพราะคู่มือเขียน `โค้ด` ไว้ในตัวหนา
    อยู่หลายที่ ถ้าไม่วนซ้ำ backtick จะโผล่ออกมาเป็นตัวอักษรจริงในเอกสาร
    """
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            add_inline(paragraph, part[2:-2], size=size, color=color, bold=True, italic=italic)
        elif part.startswith("`") and part.endswith("`"):
            run = style_run(
                paragraph.add_run(unescape(part[1:-1])),
                font=MONO,
                size=size - 1.5,
                bold=bold,
                color=NAVY,
            )
            _set(run._element.get_or_add_rPr(), "w:shd", **{"w:val": "clear", "w:fill": CODEBG})
        elif part.startswith("<http") and part.endswith(">"):
            # autolink ของ Markdown — ในเอกสารให้เห็นแค่ที่อยู่ ไม่ต้องมีวงเล็บมุมติดมา
            style_run(paragraph.add_run(part[1:-1]), size=size, bold=bold, color=NAVY)
        elif part.startswith("[") and "](" in part:
            label, target = part[1:-1].split("](", 1)
            # ลิงก์ข้ามเล่มชี้ไปไฟล์ .md — ในฉบับ .docx ให้เขียนเป็นชื่อเล่มเฉย ๆ
            style_run(
                paragraph.add_run(unescape(label)),
                size=size,
                bold=bold,
                italic=italic,
                color=NAVY if target.endswith(".md") else color,
            )
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            add_inline(paragraph, part[1:-1], size=size, color=color, bold=bold, italic=True)
        else:
            add_text(paragraph, unescape(part), size=size, bold=bold, italic=italic, color=color)


# ---------------------------------------------------------------- ตัวเขียนบล็อก


class Writer:
    def __init__(self, doc):
        self.doc = doc
        # มีอะไรถูกเขียนลงหน้าปัจจุบันแล้วหรือยัง — ใช้ตัดสินว่าหัวข้อระดับบนใบถัดไป
        # ต้องสั่งขึ้นหน้าใหม่จริงไหม ถ้าหน้ายังว่าง (เพิ่งขึ้นหน้าใหม่จากสารบัญ)
        # การสั่งซ้ำจะได้หน้าเปล่าคั่นมาหนึ่งหน้า
        self.dirty = False

    def para(self, text, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=8, size=BODY_PT, color=INK):
        p = self.doc.add_paragraph()
        p.alignment = align
        spacing(p, after=after)
        add_inline(p, text, size=size, color=color)
        self.dirty = True
        return p

    def heading(self, text, level):
        sizes = HEADING_PT
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        spacing(p, before=18 if level == 1 else 14, after=6, line=1.25)
        keep_with_next(p)
        text = re.sub(r"^(\d+)\.\s+", r"\1. ", text)
        add_inline(p, text, size=sizes[level], color=NAVY, bold=True)
        for run in p.runs:
            style_run(run, size=sizes[level], bold=True, color=NAVY)
        if level == 1:
            borders(p, bottom=RULE)
            if self.dirty:
                _set(p._p.get_or_add_pPr(), "w:pageBreakBefore", **{"w:val": "1"})
        self.dirty = True
        return p

    def bullet(self, text, *, ordered=False, index=1, indent=0.0, checkbox=False):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        spacing(p, after=4)
        pf = p.paragraph_format
        pf.left_indent = Cm(0.85 + indent)
        pf.first_line_indent = Cm(-0.55)
        marker = "☐  " if checkbox else (f"{index}.  " if ordered else "•  ")
        add_text(p, marker, color=CORAL if not checkbox else MUTED, bold=not checkbox)
        add_inline(p, text)
        self.dirty = True
        return p

    def quote(self, lines):
        p = self.doc.add_paragraph()
        # กล่องคำเตือนส่วนใหญ่ยาวแค่สองสามบรรทัด พอจัดชิดสองข้างแล้วช่องว่างเดียว
        # ที่มีในบรรทัดถูกดึงจนเป็นรูโหว่กลางประโยค — เนื้อความปกติยาวพอจึงไม่เป็นแบบนี้
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        spacing(p, before=6, after=10)
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.right_indent = Cm(0.2)
        borders(p, left="E5775A")
        shade(p, "FFF7F4")
        add_inline(p, " ".join(lines))
        self.dirty = True
        return p

    def code(self, lines):
        for i, line in enumerate(lines):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            spacing(p, before=6 if i == 0 else 0, after=6 if i == len(lines) - 1 else 0, line=1.12)
            p.paragraph_format.left_indent = Cm(0.3)
            shade(p, CODEBG)
            style_run(p.add_run(line or " "), font=MONO, size=CODE_PT, color=INK)
        self.dirty = True

    def rule(self):
        p = self.doc.add_paragraph()
        spacing(p, before=4, after=10, line=1.0)
        borders(p, bottom=RULE)
        self.dirty = True

    def image(self, path, caption):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        spacing(p, before=8, after=2, line=1.0)
        keep_with_next(p)
        borders(p, box=RULE)
        from PIL import Image

        with Image.open(path) as im:
            w, h = im.size
        width_cm = min(CONTENT_CM, 16.0 if w >= h else 11.0)
        p.add_run().add_picture(str(path), width=Cm(width_cm))
        if caption:
            cap = self.doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            spacing(cap, before=0, after=12, line=1.2)
            add_inline(cap, caption, size=CAPTION_PT, color=MUTED, italic=True)
        self.dirty = True

    def table(self, rows):
        header, body = rows[0], rows[1:]
        t = self.doc.add_table(rows=0, cols=len(header))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = False
        table_borders(t)
        widths = column_widths(rows)
        apply_grid(t, widths)
        for idx, cells in enumerate([header] + body):
            row = t.add_row()
            keep_row_together(row)
            if idx == 0:
                repeat_header(row)
            for ci, cell_text in enumerate(cells[: len(header)]):
                cell = row.cells[ci]
                cell.width = Cm(widths[ci])
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                spacing(p, before=3, after=3, line=1.25)
                add_inline(
                    p,
                    cell_text,
                    size=TABLE_PT,
                    color=RGBColor(0xFF, 0xFF, 0xFF) if idx == 0 else INK,
                    bold=idx == 0,
                )
                if idx == 0:
                    for run in p.runs:
                        style_run(run, size=TABLE_PT, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
                shade_cell(cell, "192768" if idx == 0 else (BAND if idx % 2 == 0 else "FFFFFF"))
        after = self.doc.add_paragraph()
        spacing(after, before=0, after=8, line=1.0)
        self.dirty = True
        return t


def table_borders(table):
    """เส้นตารางสีเทาอ่อน — เส้นดำของ Table Grid แข่งกับหัวตารางสีกรมท่าจนอ่านยาก"""
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        _set(tblBorders, f"w:{side}", **{"w:val": "single", "w:sz": "4", "w:space": "0", "w:color": RULE})
    tblPr.append(tblBorders)
    _set(tblPr, "w:tblLayout", **{"w:type": "fixed"})
    _set(tblPr, "w:tblW", **{"w:w": str(int(CONTENT_CM * 567)), "w:type": "dxa"})


def apply_grid(table, widths_cm):
    """
    เขียนความกว้างลง w:tblGrid ด้วย ไม่ใช่แค่ tcW ของแต่ละช่อง

    LibreOffice (และ Word เวลาเปิดไฟล์ที่ layout เป็น fixed) อ่านความกว้างจาก tblGrid
    ก่อน ถ้าตั้งแต่ tcW ตารางจะออกมาเท่ากันทุกคอลัมน์เหมือนไม่ได้ตั้งอะไรเลย
    """
    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(1, grid)
    for col in list(grid):
        grid.remove(col)
    for w in widths_cm:
        _set(grid, "w:gridCol", **{"w:w": str(int(w * 567))})


def repeat_header(row):
    """ตารางที่ยาวข้ามหน้า ต้องพาหัวตารางไปด้วย ไม่งั้นหน้าถัดไปอ่านไม่รู้เรื่อง"""
    _set(row._tr.get_or_add_trPr(), "w:tblHeader", **{"w:val": "1"})


def keep_row_together(row):
    """
    ห้ามแบ่งแถวข้ามหน้า

    ค่าตั้งต้นยอมให้แถวเดียวถูกผ่าครึ่ง ผลคือหน้าถัดไปขึ้นต้นด้วยแถวที่มีแต่คำท้าย ๆ
    ของช่องสุดท้าย ("คุณ" ลอยอยู่ช่องเดียว) อ่านแล้วไม่รู้ว่าเป็นแถวของอะไร
    """
    _set(row._tr.get_or_add_trPr(), "w:cantSplit", **{"w:val": "1"})


def column_widths(rows):
    """
    แบ่งความกว้างคอลัมน์ตามความยาวเนื้อหาจริง แล้วบีบให้อยู่ในช่วงที่อ่านได้

    ปล่อยให้ Word คำนวณเองแล้วคอลัมน์แรกมักกว้างเกินจริงจนคอลัมน์คำอธิบายแคบ
    และข้อความหักบรรทัดถี่ยิบ
    """
    cols = len(rows[0])
    weight = []
    for ci in range(cols):
        lengths = [len(strip_marks(r[ci])) for r in rows if ci < len(r)]
        weight.append(max(sum(lengths) / max(len(lengths), 1), 6))
    lo, hi = 0.13, 0.52
    share = [w / sum(weight) for w in weight]
    share = [min(max(x, lo), hi) for x in share]
    share = [x / sum(share) for x in share]
    return [CONTENT_CM * x for x in share]


def strip_marks(text):
    return re.sub(r"[*`\[\]]|\([^)]*\)", "", text)


def shade_cell(cell, fill):
    _set(cell._tc.get_or_add_tcPr(), "w:shd", **{"w:val": "clear", "w:fill": fill})


# ---------------------------------------------------------------- ตัวอ่าน Markdown

IMG = re.compile(r"^!\[(?P<alt>[^\]]*)\]\((?P<src>[^)]+)\)\s*$")
TABLE_SEP = re.compile(r"^\|[\s:|-]+\|$")


def render(md_path, writer):
    lines = md_path.read_text(encoding="utf-8").split("\n")
    i, n = 0, len(lines)
    ordered_index = 0

    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("```"):
            block, i = [], i + 1
            while i < n and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            writer.code(block)
            i += 1
            continue

        if stripped in ("---", "***", "___"):
            # ทุกหัวข้อระดับบนขึ้นหน้าใหม่แล้ว เส้นคั่นที่นำหน้ามันจะกลายเป็น
            # ขีดลอย ๆ ค้างท้ายหน้าก่อน — ข้ามไป
            nxt = next((l for l in lines[i + 1 :] if l.strip()), "")
            if not nxt.startswith("## "):
                writer.rule()
            i += 1
            continue

        m = IMG.match(stripped)
        if m:
            src = (md_path.parent / m.group("src")).resolve()
            writer.image(src, m.group("alt"))
            i += 1
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            if level == 1:  # ชื่อเล่มอยู่บนหน้าปกแล้ว
                i += 1
                continue
            writer.heading(stripped[level:].strip(), min(level - 1, 3))
            ordered_index = 0
            continue_i = i + 1
            i = continue_i
            continue

        if stripped.startswith(">"):
            block = []
            while i < n and lines[i].strip().startswith(">"):
                block.append(lines[i].strip().lstrip(">").strip())
                i += 1
            writer.quote([b for b in block if b])
            continue

        if stripped.startswith("|"):
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                raw = lines[i].strip()
                if not TABLE_SEP.match(raw):
                    rows.append([c.strip() for c in raw.strip("|").split("|")])
                i += 1
            if rows:
                writer.table(rows)
            continue

        bullet = re.match(r"^(\s*)([-*])\s+(.*)$", line)
        number = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
        if bullet or number:
            m2 = bullet or number
            indent_cm = len(m2.group(1)) / 2 * 0.4
            text = m2.group(3)
            checkbox = False
            if text.startswith("[ ] ") or text.startswith("[x] "):
                checkbox, text = True, text[4:]
            i += 1
            # บรรทัดต่อเนื่องของหัวข้อเดิม (ย่อหน้าเข้ามา และไม่ได้ขึ้นต้นหัวข้อใหม่)
            while (
                i < n
                and lines[i].strip()
                and lines[i].startswith((" ", "\t"))
                and not re.match(r"^\s*([-*]|\d+\.)\s", lines[i])
            ):
                text += " " + lines[i].strip()
                i += 1
            if number:
                ordered_index += 1
            writer.bullet(
                text,
                ordered=bool(number),
                index=ordered_index,
                indent=indent_cm,
                checkbox=checkbox,
            )
            continue

        # ย่อหน้าธรรมดา — รวมบรรทัดที่ห่อไว้ให้เป็นย่อหน้าเดียว
        block = [stripped]
        i += 1
        while (
            i < n
            and lines[i].strip()
            and not lines[i].strip().startswith(("#", ">", "|", "```", "- ", "* ", "!["))
            and not re.match(r"^\s*\d+\.\s", lines[i])
            and lines[i].strip() not in ("---", "***", "___")
        ):
            block.append(lines[i].strip())
            i += 1
        text = " ".join(block)
        if text.startswith("*") and text.endswith("*") and text.count("*") == 2:
            p = writer.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            spacing(p, before=0, after=12)
            add_inline(p, text, size=CAPTION_PT, color=MUTED)
        else:
            writer.para(text)


# ---------------------------------------------------------------- โครงเอกสาร


def setup(doc, title, subtitle):
    section = doc.sections[0]
    section.page_width, section.page_height = Cm(21.0), Cm(29.7)
    section.top_margin, section.bottom_margin = Cm(2.0), Cm(2.0)
    section.left_margin, section.right_margin = Cm(2.2), Cm(2.0)
    section.header_distance, section.footer_distance = Cm(1.1), Cm(1.1)

    normal = doc.styles["Normal"]
    normal.font.name = BODY
    normal.font.size = Pt(BODY_PT)
    rpr = normal.element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(attr), BODY)
    if rf.getparent() is None:
        rpr.insert(0, rf)
    _set(rpr, "w:szCs", **{"w:val": str(BODY_PT * 2)})

    # หน้าปกไม่ต้องมีหัว/ท้ายกระดาษ — เลขหน้า "1" บนปกดูเหมือนพิมพ์ผิด
    section.different_first_page_header_footer = True

    # หัวกระดาษ / ท้ายกระดาษ
    head = section.header.paragraphs[0]
    head.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    spacing(head, after=0, line=1.0)
    style_run(head.add_run(f"{title} · {subtitle}"), size=CHROME_PT, color=MUTED)
    borders(head, bottom=RULE)

    foot = section.footer.paragraphs[0]
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spacing(foot, before=0, after=0, line=1.0)
    style_run(foot.add_run("สถาบันข้อมูลขนาดใหญ่ (องค์การมหาชน)  ·  หน้า "), size=CHROME_PT, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run = OxmlElement("w:r")
    fld.append(run)
    foot._p.append(fld)


def cover(doc, title, subtitle, source_name, audience, updated):
    for _ in range(5):
        spacing(doc.add_paragraph(), after=0, line=1.0)
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        spacing(p, after=26, line=1.0)
        p.add_run().add_picture(str(LOGO), width=Cm(7.2))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spacing(p, after=2, line=1.2)
    style_run(p.add_run("Government Datahub Platform"), size=13, color=CORAL, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spacing(p, after=4, line=1.2)
    style_run(p.add_run(title), size=32, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spacing(p, after=30, line=1.3)
    style_run(p.add_run(subtitle), size=17, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spacing(p, after=0, line=1.5)
    style_run(p.add_run(audience), size=13, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spacing(p, after=0, line=1.5)
    style_run(p.add_run(f"ปรับปรุงล่าสุด {updated}"), size=13, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spacing(p, after=0, line=1.5)
    style_run(p.add_run(f"ฉบับ Markdown: docs/{source_name}"), size=11, color=MUTED)

    doc.add_page_break()


def toc(doc, md_path):
    heads = [
        l.strip()[3:].strip()
        for l in md_path.read_text(encoding="utf-8").split("\n")
        if l.startswith("## ")
    ]
    if len(heads) < 3:
        return
    w = Writer(doc)
    w.heading("สารบัญ", 1)
    for h in heads:
        p = doc.add_paragraph()
        spacing(p, after=3, line=1.35)
        p.paragraph_format.left_indent = Cm(0.5)
        add_text(p, "▸  ", size=BODY_PT, color=CORAL)
        add_inline(p, h)
    doc.add_page_break()


# ---------------------------------------------------------------- ฝังฟอนต์


def obfuscate(data, guid):
    """ECMA-376 §15.2.13 — 32 ไบต์แรกถูก XOR ด้วย GUID เรียงกลับหลัง"""
    key = bytes.fromhex(guid.replace("-", "").strip("{}"))
    out = bytearray(data)
    for i in range(32):
        out[i] ^= key[15 - (i % 16)]
    return bytes(out)


FACES = [
    ("embedRegular", "Sarabun-Regular.ttf"),
    ("embedBold", "Sarabun-Bold.ttf"),
    ("embedItalic", "Sarabun-Italic.ttf"),
    ("embedBoldItalic", "Sarabun-BoldItalic.ttf"),
]


def embed_fonts(docx_path, ttf_dir):
    """ยัดฟอนต์เข้าไปในไฟล์ .docx เพื่อให้เครื่องที่ไม่มี Sarabun ก็ยังเห็นถูกต้อง"""
    tmp = docx_path.with_suffix(".tmp.docx")
    parts = {}
    with zipfile.ZipFile(docx_path) as zin:
        names = zin.namelist()
        for name in names:
            parts[name] = zin.read(name)

    rels, fonts, fontlist = [], {}, []
    for idx, (tag, filename) in enumerate(FACES, start=1):
        guid = "{%s}" % str(uuid.uuid4()).upper()
        data = (ttf_dir / filename).read_bytes()
        target = f"fonts/font{idx}.odttf"
        fonts[f"word/{target}"] = obfuscate(data, guid)
        rid = f"rIdFont{idx}"
        rels.append(f'<Relationship Id="{rid}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/font" Target="{target}"/>')
        fontlist.append(f'<w:{tag} r:id="{rid}" w:fontKey="{guid}" w:subsetted="0"/>')

    font_table = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        '<w:fonts xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        f'<w:font w:name="{BODY}">'
        '<w:charset w:val="00"/><w:family w:val="swiss"/><w:pitch w:val="variable"/>'
        + "".join(fontlist)
        + "</w:font>"
        f'<w:font w:name="{MONO}"><w:charset w:val="00"/><w:family w:val="modern"/>'
        '<w:pitch w:val="fixed"/></w:font>'
        f'<w:font w:name="{SYMBOL_FONT}"><w:charset w:val="00"/><w:family w:val="swiss"/>'
        '<w:pitch w:val="variable"/></w:font>'
        "</w:fonts>"
    ).encode()

    font_rels = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        + "".join(rels)
        + "</Relationships>"
    ).encode()

    parts["word/fontTable.xml"] = font_table
    parts["word/_rels/fontTable.xml.rels"] = font_rels
    parts.update(fonts)

    # settings.xml ต้องประกาศว่าเอกสารนี้ฝังฟอนต์ไว้
    settings = parts.get("word/settings.xml", b"").decode()
    if "embedTrueTypeFonts" not in settings:
        settings = settings.replace("<w:settings", "<w:settings", 1)
        anchor = settings.index(">", settings.index("<w:settings")) + 1
        settings = settings[:anchor] + "<w:embedTrueTypeFonts/><w:saveSubsetFonts w:val=\"false\"/>" + settings[anchor:]
        parts["word/settings.xml"] = settings.encode()

    ct = parts["[Content_Types].xml"].decode()
    if 'Extension="odttf"' not in ct:
        ct = ct.replace(
            "<Types ",
            "<Types ",
            1,
        )
        anchor = ct.index(">", ct.index("<Types")) + 1
        ct = (
            ct[:anchor]
            + '<Default Extension="odttf" ContentType="application/vnd.openxmlformats-officedocument.obfuscatedFont"/>'
            + ct[anchor:]
        )
        parts["[Content_Types].xml"] = ct.encode()

    # fontTable.xml ต้องถูกอ้างจาก document.xml.rels
    drels = parts["word/_rels/document.xml.rels"].decode()
    if "fontTable.xml" not in drels:
        anchor = drels.rindex("</Relationships>")
        drels = (
            drels[:anchor]
            + '<Relationship Id="rIdFontTable" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/fontTable" Target="fontTable.xml"/>'
            + drels[anchor:]
        )
        parts["word/_rels/document.xml.rels"] = drels.encode()

    if 'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.fontTable+xml"' not in parts["[Content_Types].xml"].decode():
        ct = parts["[Content_Types].xml"].decode()
        anchor = ct.rindex("</Types>")
        ct = (
            ct[:anchor]
            + '<Override PartName="/word/fontTable.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.fontTable+xml"/>'
            + ct[anchor:]
        )
        parts["[Content_Types].xml"] = ct.encode()

    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in parts.items():
            zout.writestr(name, data)
    tmp.replace(docx_path)


# ---------------------------------------------------------------- main


def build(md_name, title, subtitle, audience, updated, ttf_dir):
    md_path = DOCS / md_name
    doc = Document()
    setup(doc, title, subtitle)
    cover(doc, title, subtitle, md_name, audience, updated)
    toc(doc, md_path)
    render(md_path, Writer(doc))

    out = OUT / (md_name.replace(".md", "") + ".docx")
    doc.save(out)
    embed_fonts(out, ttf_dir)
    return out


def main():
    ttf_dir = Path(sys.argv[1]) if len(sys.argv) > 1 else None
    if ttf_dir is None or not ttf_dir.exists():
        print("ต้องระบุโฟลเดอร์ที่แตกไฟล์ Sarabun.zip ไว้", file=sys.stderr)
        print(f"เช่น  unzip -d /tmp/sarabun {FONT_ZIP}", file=sys.stderr)
        return 1
    OUT.mkdir(parents=True, exist_ok=True)
    for md_name, title, subtitle, audience, updated in BOOKS:
        out = build(md_name, title, subtitle, audience, updated, ttf_dir)
        print(f"  {out.relative_to(ROOT)}  ({out.stat().st_size / 1024 / 1024:.1f} MB)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
